#!/usr/bin/env python3
"""Generate a professionally formatted MS Word document for the NEC Wisdom site analysis."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Definitions ───────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

NEC_BLUE = RGBColor(0x00, 0x2B, 0x5C)  # NEC corporate dark blue
NEC_LIGHT_BLUE = RGBColor(0x00, 0x7C, 0xBA)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = NEC_BLUE

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# Helper: shaded cell
DARK_BLUE_BG = '002B5C'
LIGHT_GREY = 'F2F2F2'
MEDIUM_GREY = 'E8E8E8'

def set_cell_bg(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with styled header row and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, DARK_BLUE_BG)
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_bg(cell, LIGHT_GREY)
            set_cell_text(cell, val, size=Pt(9))

    # Column widths
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Inches(w)
    return table

def add_screenshot(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(8)
        cap.runs[0].font.italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f"[Screenshot not available: {path}]")

# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('NEC Wisdom')
run.font.size = Pt(36)
run.font.color.rgb = NEC_BLUE
run.bold = True
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Website Analysis Report')
run.font.size = Pt(24)
run.font.color.rgb = NEC_LIGHT_BLUE
run.font.name = 'Calibri'

doc.add_paragraph()

url_p = doc.add_paragraph()
url_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = url_p.add_run('https://wisdom.nec.com/')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()
doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Comprehensive site analysis covering templates, components,\n'
                    'integrations, page inventory, and migration complexity assessment\n'
                    'for Adobe Edge Delivery Services migration.')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('March 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Site Overview',
    '2. Templates Inventory',
    '3. Blocks & Components Catalog',
    '4. Page Counts by Template',
    '5. Third-Party Integrations & Embedded Services',
    '6. Enterprise System Integrations (REST/API Calls)',
    '7. Complex Use Cases & Observations',
    '8. Migration Estimates',
    '9. Screenshots'
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. SITE OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Site Overview', level=1)

doc.add_paragraph(
    'NEC Wisdom (wisdom.nec.com) is a Japanese-language business and technology content media '
    'platform operated by NEC Corporation. Branded as "Business Leaders Square wisdom," it serves '
    'as a thought leadership hub delivering articles, expert columns, special features, and event '
    'coverage focused on digital transformation, AI, cybersecurity, smart cities, and other '
    'technology-driven business topics.'
)
doc.add_paragraph(
    'The site targets Japanese business executives, IT decision-makers, and technology professionals. '
    'Content is organized across multiple taxonomies: themes (technology topics), industries, '
    'special features (curated collections), and author series (regular columns by named experts).'
)

doc.add_heading('Key Characteristics', level=2)
overview_data = [
    ['URL', 'https://wisdom.nec.com/'],
    ['Language', 'Japanese (ja)'],
    ['Platform', 'Custom CMS (server-rendered HTML with jQuery)'],
    ['Site Type', 'B2B Content Media / Thought Leadership Portal'],
    ['Total Indexed Pages', '~1,362 (from sitemap)'],
    ['Content Sections', 'Feature (themes), Business, Series (columns), Special, Events, Technology, Innovation'],
    ['Member System', 'My NEC registration (auth.nec.com) for event access and downloads'],
    ['Primary Audience', 'Japanese business leaders, IT decision-makers, technology professionals'],
]
add_styled_table(doc, ['Attribute', 'Details'], overview_data, col_widths=[2.0, 4.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. TEMPLATES INVENTORY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Templates Inventory', level=1)

doc.add_paragraph(
    'The site uses a consistent design system with a limited number of distinct page templates. '
    'All templates share a common header (logo, search, member registration, NEC corporate link), '
    'global navigation (5-item mega menu: Themes, Industries, Special Features, Series, Events), '
    'a member registration CTA section, cookie consent banner, and a comprehensive mega-footer.'
)

templates = [
    ['T1', 'Homepage', '/',
     'Hero carousel with featured articles, category cards grid (Themes, Industries, Special, Series), '
     'latest articles feed, ranking sidebar (weekly/monthly tabs), member CTA',
     '1'],
    ['T2', 'Article Detail', '/ja/feature/ai/2025090501/index.html',
     'Breadcrumb, article header (title, date, reading time, share button), collapsible summary/TOC, '
     'rich body content (headings, images, comparison tables, figures, lists), related articles "OTHER" section, '
     'related services section, latest articles sidebar, ranking sidebar',
     '~1,100'],
    ['T3', 'Theme/Topic Listing', '/ja/feature/ai/index.html',
     'Hero banner with topic description, H1 title, explanatory paragraphs, '
     'article card grid (thumbnail, title, date, reading time, tags), "Load More" button, '
     '"OTHER" section with related theme cards',
     '~16'],
    ['T4', 'Special Feature Collection', '/ja/special/digital_ethics/index.html',
     'Hero banner, curated introduction text, sectioned article lists (e.g., "Expert Interviews", '
     '"Survey Data Analysis"), article cards with dates and tags',
     '~11'],
    ['T5', 'Series/Author Column', '/ja/series/orita/index.html',
     'Author hero section, series description text, article card grid, '
     'author biography section with photo, "Load More" button',
     '~7'],
    ['T6', 'Event Hub', '/ja/event/index.html',
     'Featured event hero (with status badge), event report sections with grouped articles, '
     'past events grid, "OTHER" category navigation cards, "Load More" button',
     '1'],
    ['T7', 'Event Detail', '/ja/event/seminer/20260108/index.html',
     'Event header with title and share button, event overview panel (dates, fee), '
     'CTA registration button, embedded video player (Brightcove), session schedule with '
     'speaker profiles (expandable), tags, latest articles sidebar, ranking sidebar',
     '~50'],
    ['T8', 'Article Filter/Search', '/ja/article.html?B02',
     'Dynamic filtering page using URL parameters for theme/industry codes, '
     'article card grid loaded via JavaScript/AJAX, tag-based navigation',
     '~27'],
    ['T9', 'Contact Form (CGI)', '/cgi-bin/ja/contact/index.cgi',
     'Legacy CGI-based contact form, hosted under /cgi-bin/ path',
     '1'],
    ['T10', 'Utility Pages', '/ja/wisdom/index.html, /ja/privacy.html, /ja/termofuse.html, /ja/sitemap.html',
     'About page, privacy policy, terms of use, sitemap - static content pages',
     '~4'],
]

add_styled_table(doc,
    ['ID', 'Template Name', 'Example URL', 'Key Components', 'Est. Pages'],
    templates,
    col_widths=[0.4, 1.2, 1.5, 2.6, 0.6])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. BLOCKS & COMPONENTS CATALOG
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Blocks & Components Catalog', level=1)

doc.add_paragraph(
    'The following reusable UI components appear across multiple templates. '
    'These would need to be mapped to EDS blocks or implemented as custom blocks during migration.'
)

doc.add_heading('3.1 Global Components (All Pages)', level=2)
global_blocks = [
    ['Global Header', 'Fixed top bar: Wisdom logo, member registration link, search button, NEC corporate logo/link',
     'All pages', 'Medium - Mega menu with 5 dropdown panels'],
    ['Global Navigation', '5-item mega menu: Themes (テーマ), Industries (業種), Special (特集), Series (連載), Events (イベント). Themes sub-nav has 3 groups: Technology, Urban/Social, Business/Economy',
     'All pages', 'High - Complex multi-level dropdown menus'],
    ['Breadcrumb', 'Hierarchical breadcrumb trail (Top > Section > Subsection > Page)',
     'All pages except homepage', 'Low'],
    ['Cookie Consent Banner', 'Bottom banner with privacy policy link and OK button',
     'All pages', 'Low'],
    ['Member Registration CTA', 'Full-width section with "MEMBER" heading, description text, and dual CTA links (My NEC registration, My NEC details)',
     'All pages', 'Low'],
    ['Mega Footer', 'Multi-column footer: Theme links (3 subcategories), Industry links (13 items), Special feature links (11 items), Series links (7 authors), utility links (Events, About, Sitemap, Privacy, Terms, Contact), copyright',
     'All pages', 'Medium - Large link taxonomy'],
]
add_styled_table(doc,
    ['Component', 'Description', 'Used On', 'Migration Complexity'],
    global_blocks,
    col_widths=[1.2, 2.8, 1.0, 1.3])

doc.add_heading('3.2 Content Components', level=2)
content_blocks = [
    ['Hero Carousel', 'Full-width image carousel with article card overlays on homepage. Auto-rotating with indicators.',
     'Homepage', 'High'],
    ['Article Card (Standard)', 'Thumbnail image, title (H3 link), date, reading time, tag pills (theme + industry)',
     'Theme listing, Special, Series, Event reports', 'Low'],
    ['Article Card (Large)', 'Wide card with larger image, used in featured positions on event pages',
     'Event hub featured section', 'Low'],
    ['Category Card Grid', '4-column grid of category cards with thumbnails and labels (e.g., Theme subcategories)',
     'Homepage, Theme listing "OTHER" section', 'Low'],
    ['Article Summary/TOC', 'Collapsible accordion containing article summary and table of contents with anchor links',
     'Article detail pages', 'Medium'],
    ['Rich Article Body', 'Long-form content with H2/H3 headings, paragraphs, inline images, comparison tables, figures with captions, blockquotes, numbered/bulleted lists',
     'Article detail pages', 'Medium - Complex nested content'],
    ['Comparison Table', 'Styled data table within article content (e.g., feature comparison charts)',
     'Article detail pages', 'Medium'],
    ['Related Articles ("OTHER")', 'Section with H1 heading, H2 subcategory, horizontal scrollable article cards',
     'Theme listing, Article detail', 'Medium'],
    ['Related Services', 'Section listing NEC services related to the article topic',
     'Article detail pages', 'Low'],
    ['Share Button', 'Social sharing dropdown button (SNS share)',
     'Article detail, Event detail', 'Low'],
    ['Tag Pills', 'Inline tag links with theme and industry classification codes',
     'Article cards, Article detail pages', 'Low'],
    ['Load More Button', 'AJAX-powered "もっと見る" button for progressive article loading',
     'All listing pages', 'Medium - Requires JS implementation'],
]
add_styled_table(doc,
    ['Component', 'Description', 'Used On', 'Migration Complexity'],
    content_blocks,
    col_widths=[1.3, 2.5, 1.2, 1.3])

doc.add_heading('3.3 Sidebar Components', level=2)
sidebar_blocks = [
    ['Latest Articles Sidebar', 'Vertical list of 5-6 latest articles with thumbnails, theme label, and title. Full-width cards.',
     'Article detail, Event detail', 'Low'],
    ['Ranking Sidebar', 'Tabbed component (Weekly/Monthly) showing numbered list of top 5 popular articles by title',
     'Article detail, Event detail, Homepage', 'Medium - Tab switching + analytics data'],
    ['Author Biography', 'Author photo, name, and biographical paragraph displayed in sidebar',
     'Series listing pages', 'Low'],
]
add_styled_table(doc,
    ['Component', 'Description', 'Used On', 'Migration Complexity'],
    sidebar_blocks,
    col_widths=[1.3, 2.8, 1.2, 1.0])

doc.add_heading('3.4 Event-Specific Components', level=2)
event_blocks = [
    ['Event Overview Panel', 'Styled panel with event dates, fee, and registration CTA button. Blue header with "EVENT OVERVIEW".',
     'Event detail pages', 'Medium'],
    ['Video Player', 'Embedded Brightcove video player for event recordings/streams',
     'Event detail pages', 'High - Third-party embed'],
    ['Session Schedule', 'Time-slotted session list with titles, descriptions, and expandable speaker profiles with photos',
     'Event detail pages', 'Medium'],
    ['Speaker Profile Card', 'Expandable accordion with speaker photo, name, title, and affiliation',
     'Event detail pages', 'Low'],
    ['Event Status Badge', 'Badge overlay (e.g., "アーカイブ配信中" / Archive Streaming, "申込締切" / Applications Closed)',
     'Event listing, Event detail', 'Low'],
    ['Event Registration CTA', 'Large button linking to auth.nec.com registration/seminar forms',
     'Event detail pages', 'Low'],
]
add_styled_table(doc,
    ['Component', 'Description', 'Used On', 'Migration Complexity'],
    event_blocks,
    col_widths=[1.3, 2.8, 1.2, 1.0])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. PAGE COUNTS BY TEMPLATE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Page Counts by Template', level=1)

doc.add_paragraph(
    'The sitemap (wisdom.nec.com/ja/sitemap.html) contains 1,362 unique URLs. '
    'Pages are organized under the /ja/ path prefix with the following URL structure and distribution:'
)

doc.add_heading('4.1 Page Count by URL Section', level=2)
section_counts = [
    ['/ja/feature/', 'Theme/Topic articles', '467', '34.3%'],
    ['/ja/series/', 'Author column articles', '241', '17.7%'],
    ['/ja/business/', 'Business case studies', '160', '11.7%'],
    ['/ja/events/ + /ja/event/', 'Event pages', '152 (94+58)', '11.2%'],
    ['/ja/article/', 'Article filter pages', '84', '6.2%'],
    ['/ja/collaboration/', 'Collaboration articles', '70', '5.1%'],
    ['/ja/technology/', 'Technology articles', '68', '5.0%'],
    ['/ja/innovation/', 'Innovation articles', '41', '3.0%'],
    ['/ja/special/', 'Special features', '36', '2.6%'],
    ['/ja/solutions/', 'Solution pages', '16', '1.2%'],
    ['/ja/strategy/', 'Strategy articles', '6', '0.4%'],
    ['Other (utility, etc.)', 'Sitemap, privacy, terms, about, etc.', '~21', '1.5%'],
]
add_styled_table(doc,
    ['URL Pattern', 'Content Type', 'Page Count', '% of Total'],
    section_counts,
    col_widths=[1.5, 2.0, 1.0, 0.8])

doc.add_heading('4.2 Estimated Template Distribution', level=2)
template_dist = [
    ['T2 - Article Detail', '~1,100', '80.8%', 'Bulk of content; articles across all sections (feature, business, series, technology, etc.)'],
    ['T8 - Article Filter/Search', '~84', '6.2%', 'Dynamic filter pages using query parameters'],
    ['T7 - Event Detail', '~50', '3.7%', 'Individual event/seminar/workshop pages'],
    ['T3 - Theme/Topic Listing', '~16', '1.2%', 'One per major theme (AI, 5G, cybersecurity, etc.)'],
    ['T4 - Special Feature Collection', '~11', '0.8%', 'Curated editorial collections'],
    ['T5 - Series/Author Column', '~7', '0.5%', 'One per regular columnist'],
    ['T1 - Homepage', '1', '0.1%', 'Single homepage'],
    ['T6 - Event Hub', '1', '0.1%', 'Single event listing page'],
    ['T9 - Contact Form', '1', '0.1%', 'CGI-based contact form'],
    ['T10 - Utility Pages', '~4', '0.3%', 'About, privacy, terms, sitemap'],
]
add_styled_table(doc,
    ['Template', 'Est. Pages', '% of Total', 'Notes'],
    template_dist,
    col_widths=[1.8, 0.8, 0.7, 3.0])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. THIRD-PARTY INTEGRATIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Third-Party Integrations & Embedded Services', level=1)

doc.add_paragraph(
    'NEC Wisdom employs a comprehensive analytics, advertising, and personalization stack. '
    'The following third-party services were detected through network request analysis:'
)

integrations = [
    ['Tealium IQ (Tag Management)', 'tiqcdn.com, tealiumiq.com', 'Tag management system orchestrating all other tags. Primary TMS for the site.',
     'High', 'Must replicate tag management strategy; Tealium loads many downstream tags'],
    ['Adobe DTM / Adobe Analytics', 'adobedtm.com', 'Web analytics tracking via Adobe Dynamic Tag Manager',
     'High', 'Core analytics - requires Adobe Analytics implementation on EDS'],
    ['Adobe Audience Manager', 'nec.demdex.net (iframe)', 'Audience segmentation and DMP for targeted content/ads',
     'High', 'DMP integration via iframe; requires careful migration'],
    ['Google Tag Manager', 'googletagmanager.com', 'Secondary tag manager for Google services',
     'Medium', 'Standard GTM implementation; straightforward to migrate'],
    ['Google Ads / DoubleClick', 'doubleclick.net', 'Digital advertising and remarketing pixels',
     'Medium', 'Ad tracking; standard implementation'],
    ['Facebook Pixel', 'connect.facebook.net', 'Social advertising and conversion tracking',
     'Low', 'Standard pixel; easy to migrate'],
    ['TrenDemon', 'trendemon.com, trackingapi.trendemon.com', 'Content personalization, journey orchestration, and attribution analytics',
     'High', 'Content recommendations engine; critical for content engagement strategy'],
    ['Oracle Eloqua', 'img03.en25.com', 'Marketing automation platform for lead nurturing and email campaigns',
     'High', 'Form handling and lead management integration'],
    ['Docodoco', 'docodoco.jp', 'Japanese IP geolocation service for company identification of B2B visitors',
     'Medium', 'Japan-specific B2B visitor intelligence service'],
    ['Brightcove', 'players.brightcove.net', 'Enterprise video hosting and player for event recordings',
     'Medium', 'Video embed; requires player embed code migration'],
    ['My NEC / auth.nec.com', 'auth.nec.com', 'NEC corporate single sign-on for member registration and event signup',
     'High', 'Authentication system with SSO; deep integration with event CTAs'],
]
add_styled_table(doc,
    ['Service', 'Domain(s)', 'Purpose', 'Impact', 'Migration Notes'],
    integrations,
    col_widths=[1.2, 1.2, 1.5, 0.6, 1.8])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. ENTERPRISE SYSTEM INTEGRATIONS (REST/API CALLS)
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Enterprise System Integrations (REST/API Calls)', level=1)

doc.add_paragraph(
    'A deep analysis of network traffic, JavaScript source code, and runtime objects reveals '
    'a comprehensive set of enterprise system integrations communicating via REST APIs, JSONP callbacks, '
    'and XHR/Fetch requests. No SOAP/WSDL endpoints were detected; the site exclusively uses '
    'REST-style HTTP APIs. The following table catalogs all identified API-level integrations.'
)

doc.add_heading('6.1 Internal Content & Data APIs', level=2)
doc.add_paragraph(
    'The site uses several internal API endpoints to dynamically load content, navigation fragments, '
    'and article data. These are fetched client-side via JavaScript fetch() and jQuery $.ajax() calls.'
)

internal_apis = [
    ['Content Catalog CSV', 'GET', '/ja/common/csv/articles.csv',
     'fetch() in main.js',
     'Returns 1,460-line CSV file containing the complete article catalog: titles, URLs, publish dates, '
     'og:image paths, thumbnails, theme codes (B01-B29), industry codes (C01-C14), reading time, '
     'and article type classifications. Powers the article filtering system (/ja/article.html?B02) and '
     '"Load More" pagination. Client-side JS parses and filters this CSV.'],
    ['RSS/XML Feed', 'GET ($.ajax)', '/ja/rss.xml',
     'jQuery $.ajax() with dataType: "xml"',
     'Returns 598KB RSS 2.0 XML feed with full article metadata. Used for ranking sidebar population '
     'and content syndication. Includes dc:, sy:, rdf: XML namespaces.'],
    ['Header Fragment', 'GET', '/nf_global/wisdom/v2/local/ja/include/header.html',
     'fetch()',
     'Server-side include loaded via client-side fetch. Returns HTML fragment (1.7KB) for site header '
     '(logo, member registration link, search button, NEC corporate link).'],
    ['Global Navigation', 'GET', '/nf_global/wisdom/v2/local/ja/include/gnavi.html',
     'fetch()',
     'Returns 22.8KB HTML fragment containing the full mega menu navigation structure with all '
     'dropdown panels, links, and category hierarchies.'],
    ['Member CTA Section', 'GET', '/nf_global/wisdom/v2/local/ja/include/member.html',
     'fetch()',
     'Returns 1.3KB HTML fragment for the member registration CTA section with My NEC links.'],
    ['Page-Top Button', 'GET', '/nf_global/wisdom/v2/local/ja/include/page-top.html',
     'fetch()',
     'Returns HTML fragment for the "Back to top" floating button.'],
    ['Footer Fragment', 'GET', '/nf_global/wisdom/v2/local/ja/include/footer.html',
     'fetch()',
     'Returns HTML fragment for the comprehensive mega footer with all category links.'],
]
add_styled_table(doc,
    ['Endpoint', 'Method', 'URL / Path', 'Call Mechanism', 'Description & Data'],
    internal_apis,
    col_widths=[1.0, 0.5, 1.8, 1.0, 2.0])

doc.add_heading('6.2 Third-Party REST API Integrations', level=2)
doc.add_paragraph(
    'The following external REST APIs are called from the client-side during page load and user interactions. '
    'These represent enterprise-grade integrations for personalization, analytics, visitor intelligence, '
    'and marketing automation.'
)

external_apis = [
    ['TrenDemon\nContent Personalization\nREST API (JSONP)', 'GET',
     'trackingapi.trendemon.com/api/settings/{accountId}\n'
     'trackingapi.trendemon.com/api/experience/personal-stream\n'
     'trackingapi.trendemon.com/api/experience/personal\n'
     'trackingapi.trendemon.com/api/experience/personal-embedded',
     'Account ID: 2316. Returns JSON via JSONP callback. Settings API returns 39 CTAs, '
     'multi-domain tracking (wisdom.nec.com, jpn.nec.com, eform.nec.com, mywisdom.nec.com, auth.nec.com). '
     'Personal-stream and personal APIs return personalized content recommendations. '
     'Has inbound streams enabled. Drives content recommendation widgets on article pages.'],
    ['Tealium Visitor Service\nREST API (JSONP)', 'GET',
     'visitor-service-ap-northeast-1.tealiumiq.com/nec/main/{visitorId}',
     'Returns visitor profile data via JSONP callback (visitor_service_cb_info). '
     'Used for cross-session visitor identification and audience segmentation. '
     'Region: ap-northeast-1 (Tokyo). Profile/account: nec/main.'],
    ['Tealium Event\nCollection API', 'POST (XHR)',
     'collect-ap-northeast-1.tealiumiq.com/nec/main/2/i.gif',
     'XMLHttpRequest POST for event tracking. Sends page view, click, and conversion events to '
     'Tealium\'s server-side collection endpoint. Multiple calls per page load. '
     'Tealium utag version: ut4.45 (prod build 202505120257).'],
    ['Docodoco B2B\nIP Geolocation\nREST API (v5)', 'GET',
     'api.docodoco.jp/v5/docodoco?key={apiKey}',
     'Japanese B2B visitor intelligence API. Identifies visiting company based on IP address. '
     'API version 5 with authenticated key. Returns company name, industry, size, and location data. '
     'Results feed into Tealium/analytics for B2B lead scoring.'],
    ['Oracle Eloqua\nMarketing Automation', 'GET (Script)',
     'img03.en25.com/i/elqCfg.min.js',
     'Loads Eloqua configuration and visitor tracking. Sets window objects: _elqQ, _elq, elqCookieValue. '
     'Handles form pre-population and visitor identification for marketing automation workflows. '
     'No direct form submissions observed (forms redirect to auth.nec.com).'],
    ['Adobe Analytics\n(SiteCatalyst)\nvia Adobe Launch', 'GET/POST',
     'assets.adobedtm.com/224e078edf50/2cbf430e2d3d/...',
     'Adobe Launch (Turbine v29.0.0, built 2026-02-24). Property: "NEC JPN" (PR4b7a640b49714def). '
     'Report suite: s_account="neccojp". Loads satellite library and multiple rule extensions '
     '(RC69e17d, EXdb0224e). Tracks page views, events, and custom variables.'],
    ['Google Ads\nConversion API', 'POST (Fetch)',
     'www.google.com/ccm/collect\n'
     'googleads.g.doubleclick.net/pagead/viewthroughconversion/{id}/',
     'Two Google Ads accounts: AW-11334075662 (primary) and AW-580845434 (secondary). '
     'Sends page_view and conversion events via Fetch POST. Includes full page metadata, '
     'device info, and consent state (gcd parameter). Loaded via GTM.'],
    ['Microsoft Clarity\nSession Analytics', 'POST',
     'o.clarity.ms/collect\n'
     'www.clarity.ms/tag/p0dwo13zb4',
     'Session recording and heatmap analytics. Project ID: p0dwo13zb4. '
     'Multiple POST calls per page for session replay data collection. '
     'Captures clicks, scrolls, and mouse movements.'],
    ['Yahoo Japan APM\nAdvertising Measurement', 'GET (Fetch)',
     'apm.yahoo.co.jp/rt/?p=GCRWA6ZFPR',
     'Yahoo Japan advertising performance measurement. Project ID: GCRWA6ZFPR. '
     'Tracks page views with referrer, user agent brands, and platform data. '
     'Loaded via Tealium utag extension (utag.94.js).'],
    ['Facebook Conversion API\n(Dual Pixel)', 'GET (Script)',
     'connect.facebook.net/signals/config/{pixelId}\n'
     'connect.facebook.net/en_US/fbevents.js',
     'Two Facebook Pixel IDs: 972786687120127 and 3096900423961692. '
     'Tracks PageView events. Extended match enabled with 250+ feature flags. '
     'Domain verified for wisdom.nec.com.'],
    ['NEC Authentication\nSSO (My NEC)', 'GET (Redirect)',
     'auth.nec.com/register\n'
     'auth.nec.com/seminar/{eventCode}',
     'Member registration and event signup endpoints. All links append Tealium visitor tracking '
     'parameter (teal_wdm={visitorId}). Enables cross-domain visitor stitching between '
     'wisdom.nec.com and auth.nec.com for conversion attribution.'],
]
add_styled_table(doc,
    ['Service / Protocol', 'Method', 'Endpoint URL(s)', 'Details & Data Flow'],
    external_apis,
    col_widths=[1.2, 0.5, 2.0, 2.6])

doc.add_heading('6.3 API Integration Architecture Summary', level=2)
doc.add_paragraph(
    'The site employs a layered API integration architecture orchestrated primarily through Tealium IQ '
    'as the tag management system. The data flow is as follows:'
)
doc.add_paragraph(
    '1. Page Load: Tealium IQ (utag.js) initializes and orchestrates all downstream tag loading.\n'
    '2. Visitor Identification: Tealium Visitor Service retrieves cross-session visitor profile; '
    'Docodoco identifies the visitor\'s company via IP geolocation; Eloqua sets marketing automation cookies.\n'
    '3. Content Delivery: Internal CSV API provides full article catalog for client-side filtering; '
    'SSI fragments deliver header/nav/footer components; RSS XML provides ranking data.\n'
    '4. Personalization: TrenDemon REST APIs return personalized content recommendations based on '
    'visitor behavior, journey stage, and account identification.\n'
    '5. Analytics Collection: Adobe Analytics (SiteCatalyst), Google Ads, Microsoft Clarity, '
    'Yahoo Japan APM, and Facebook Pixel all receive event data via their respective APIs.\n'
    '6. Authentication: My NEC SSO handles member registration and event signup with Tealium '
    'visitor ID passthrough for cross-domain attribution.'
)

doc.add_heading('6.4 Migration Impact Assessment', level=2)
api_impact = [
    ['Internal CSV Content API', 'High',
     'The /ja/common/csv/articles.csv endpoint is the backbone of the article filtering system. '
     'Must be replaced with EDS-compatible content index (e.g., query-index.json) or static pre-rendering.'],
    ['Server-Side Include Fragments', 'High',
     'Five HTML fragments are fetched client-side. In EDS, these become authored fragments or '
     'auto-blocks. Navigation fragment (22.8KB) is especially complex.'],
    ['RSS/XML Feed', 'Medium',
     'Used for ranking sidebar. EDS can generate query-index.json as a replacement. '
     'RSS feed may need to be maintained for external syndication.'],
    ['TrenDemon Personalization API', 'High',
     'Four REST endpoints drive content personalization. TrenDemon script must load on EDS pages; '
     'cross-domain tracking configuration spans 5 NEC domains.'],
    ['Tealium Visitor Service', 'High',
     'Core to visitor identification and all downstream analytics. Must be first script to load on EDS.'],
    ['Docodoco B2B Intelligence', 'Medium',
     'API key-authenticated endpoint. Script tag must be included in EDS head.html or delayed.js.'],
    ['Adobe Analytics / Launch', 'High',
     'Adobe Launch property must be configured for new EDS domain. Report suite "neccojp" unchanged.'],
    ['Dual Google Ads Accounts', 'Low',
     'Standard GTM container handles both accounts. GTM snippet in head.html.'],
    ['Microsoft Clarity', 'Low',
     'Single script tag inclusion. May conflict with other session recording tools.'],
    ['Yahoo Japan APM', 'Low',
     'Loaded via Tealium; will work automatically once Tealium is configured.'],
    ['Facebook Dual Pixel', 'Low',
     'Standard pixel code. Include via Tealium or direct script tag.'],
    ['My NEC Authentication', 'Medium',
     'CTA links must preserve teal_wdm tracking parameter format for cross-domain attribution.'],
]
add_styled_table(doc,
    ['API Integration', 'Migration Impact', 'Notes'],
    api_impact,
    col_widths=[1.5, 0.8, 4.0])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. COMPLEX USE CASES & OBSERVATIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Complex Use Cases & Observations', level=1)

doc.add_heading('7.1 Japanese Language & Content', level=2)
doc.add_paragraph(
    'The entire site is in Japanese with no multilingual variants. Content authoring must maintain '
    'Japanese typography standards, including proper handling of kanji, hiragana, katakana, and '
    'mixed Japanese-English text. URL structures use romanized paths (/ja/feature/ai/) while page '
    'content is entirely in Japanese. Font rendering and line-breaking rules for Japanese text '
    'require special attention in CSS.'
)

doc.add_heading('7.2 Dynamic Article Filtering System', level=2)
doc.add_paragraph(
    'The site uses a dynamic article filtering system at /ja/article.html with URL parameter codes '
    '(e.g., ?B02 for AI, ?C01 for Finance industry). This system uses JavaScript/AJAX to load filtered '
    'article lists without page reloads. There are approximately 27+ unique filter codes covering '
    'themes (B-codes) and industries (C-codes). This dynamic filtering system would need to be '
    'reimplemented as either static listing pages or a client-side filtering solution in EDS.'
)

doc.add_heading('7.3 Mega Menu Navigation', level=2)
doc.add_paragraph(
    'The global navigation features a 5-item mega menu with multi-panel dropdowns. The "Themes" menu '
    'alone has 3 subcategories (Technology, Urban/Social, Business/Economy) with 16+ links. '
    'The "Industries" menu lists 13 industry categories. The "Special Features" and "Series" menus '
    'list 11 and 7 items respectively. This complex navigation structure requires a sophisticated '
    'nav implementation in EDS.'
)

doc.add_heading('7.4 Member Registration & Event System', level=2)
doc.add_paragraph(
    'The site deeply integrates with NEC\'s My NEC authentication system (auth.nec.com). Event '
    'registration flows redirect to auth.nec.com/seminar/ URLs with tracking parameters. The member '
    'registration CTA appears on every page. This authentication dependency means EDS pages must '
    'maintain proper redirect URLs with Tealium tracking parameters (teal_wdm=...).'
)

doc.add_heading('7.5 Content Personalization (TrenDemon)', level=2)
doc.add_paragraph(
    'TrenDemon provides content personalization, recommending related articles and orchestrating '
    'visitor journeys. This includes tracking APIs (trackingapi.trendemon.com) that influence which '
    'content is displayed to returning visitors. Migration must account for maintaining personalization '
    'capabilities or replacing them with alternative EDS-compatible solutions.'
)

doc.add_heading('7.6 Legacy CGI Contact Form', level=2)
doc.add_paragraph(
    'The contact page uses a legacy CGI-bin application (/cgi-bin/ja/contact/index.cgi). This '
    'server-side form processing cannot be migrated directly to EDS and would need to be replaced '
    'with a modern form solution (e.g., Adobe Forms, third-party form service, or API endpoint).'
)

doc.add_heading('7.7 Dual Tag Management Architecture', level=2)
doc.add_paragraph(
    'The site runs both Tealium IQ AND Google Tag Manager simultaneously. Tealium appears to be '
    'the primary TMS managing Adobe Analytics, while GTM handles Google-specific services. This '
    'dual TMS architecture adds complexity to migration since both need to be properly replicated '
    'on the EDS implementation.'
)

doc.add_heading('7.8 Brightcove Video Integration', level=2)
doc.add_paragraph(
    'Event detail pages embed Brightcove video players for seminar recordings. The player uses '
    'account ID 4598493582001 and custom plugins including Tealium analytics integration. '
    'Video embeds need proper responsive container handling and script loading in EDS.'
)

doc.add_heading('7.9 B2B Visitor Intelligence', level=2)
doc.add_paragraph(
    'Docodoco, a Japan-specific IP geolocation service, identifies the companies that visitors '
    'belong to based on their IP addresses. This B2B visitor intelligence feeds into the lead '
    'generation pipeline alongside Oracle Eloqua. Maintaining this capability requires ensuring '
    'the Docodoco script loads correctly on EDS pages.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. MIGRATION ESTIMATES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Migration Estimates', level=1)

doc.add_heading('8.1 Complexity Assessment', level=2)
complexity = [
    ['Overall Site Complexity', 'Medium-High',
     'Japanese-only content media site with ~1,362 pages, complex taxonomy, heavy analytics stack, and authentication integration'],
    ['Content Volume', 'High',
     '~1,100 article pages plus ~260 listing/event/utility pages; mostly standardized templates'],
    ['Template Diversity', 'Medium',
     '10 distinct templates, but ~80% of pages use a single article detail template'],
    ['Navigation Complexity', 'High',
     'Multi-level mega menu with 50+ links across 5 categories and subcategories'],
    ['Integration Complexity', 'High',
     '11 third-party services including dual TMS (Tealium + GTM), DMP (Adobe AM), personalization (TrenDemon), and authentication (My NEC)'],
    ['Custom Functionality', 'Medium',
     'Dynamic article filtering, AJAX load-more, collapsible TOC, tabbed ranking, video embeds'],
    ['Content Authoring Impact', 'Medium',
     'Japanese text requires attention to typography; standardized article format simplifies authoring'],
]
add_styled_table(doc,
    ['Dimension', 'Rating', 'Details'],
    complexity,
    col_widths=[1.5, 1.0, 3.8])

doc.add_heading('8.2 Block Development Estimates', level=2)
block_estimates = [
    ['Global Header + Mega Menu', 'Custom', 'High', 'Complex multi-level dropdown navigation with 50+ links'],
    ['Breadcrumb', 'Standard', 'Low', 'Simple hierarchical breadcrumb; map to EDS default'],
    ['Hero Carousel', 'Custom', 'High', 'Auto-rotating carousel with article card overlays'],
    ['Article Card Grid', 'Cards variant', 'Low', 'Map to EDS Cards block with custom variant for tags/date'],
    ['Article Summary/TOC', 'Custom', 'Medium', 'Collapsible accordion with anchor link TOC'],
    ['Rich Article Body', 'Default content', 'Low', 'Standard EDS default content (headings, paragraphs, images, lists)'],
    ['Comparison Table', 'Table block', 'Low', 'EDS Table block with styling'],
    ['Related Articles Section', 'Cards variant', 'Medium', 'Horizontal scrollable card section with section header'],
    ['Latest Articles Sidebar', 'Cards variant', 'Low', 'Vertical card list for sidebar layout'],
    ['Ranking Sidebar (Tabbed)', 'Custom', 'Medium', 'Tabbed weekly/monthly ranking with numbered list; requires JS'],
    ['Member CTA Section', 'Custom', 'Low', 'Full-width CTA with text and dual buttons'],
    ['Event Overview Panel', 'Custom', 'Medium', 'Styled information panel with registration CTA'],
    ['Video Player Embed', 'Embed block', 'Medium', 'Brightcove player embed with custom styling'],
    ['Session Schedule', 'Custom', 'Medium', 'Time-slotted layout with expandable speaker profiles'],
    ['Load More / Pagination', 'Custom', 'Medium', 'AJAX-based progressive content loading'],
    ['Cookie Consent Banner', 'Custom', 'Low', 'Simple banner with OK button'],
    ['Mega Footer', 'Custom', 'Medium', 'Multi-column taxonomy footer with 60+ links'],
    ['Tag Pills', 'Custom', 'Low', 'Inline colored tag links'],
    ['Share Button', 'Custom', 'Low', 'Social sharing dropdown'],
    ['Article Filter System', 'Custom', 'High', 'Dynamic filtering with URL parameters; may need redesign for EDS'],
]
add_styled_table(doc,
    ['Block/Component', 'EDS Mapping', 'Effort', 'Notes'],
    block_estimates,
    col_widths=[1.5, 1.0, 0.6, 3.2])

doc.add_heading('8.3 Migration Effort Summary', level=2)
effort = [
    ['Design System & Global Styles', 'Extract NEC Wisdom design tokens (colors, typography, spacing) to CSS custom properties. Dark header/footer theme, Japanese font stack.',
     'Medium'],
    ['Global Header + Navigation', 'Multi-level mega menu with 5 categories, search, member link. Most complex single component.',
     'High'],
    ['Article Detail Template (T2)', 'Primary template covering ~80% of pages. Rich content with sidebar, TOC, related articles.',
     'Medium'],
    ['Listing Templates (T3-T6)', 'Four listing variations (Theme, Special, Series, Event hub). Share common card grid pattern with "Load More."',
     'Medium'],
    ['Event Detail Template (T7)', 'Video player, session schedule, speaker profiles, registration CTAs.',
     'Medium-High'],
    ['Article Filter System (T8)', 'Dynamic filtering using URL params. Requires JS solution or static pre-generation.',
     'High'],
    ['Homepage (T1)', 'Hero carousel, category grids, latest articles, ranking sidebar.',
     'Medium-High'],
    ['Analytics & Tag Management', 'Dual TMS (Tealium + GTM), Adobe Analytics, Adobe Audience Manager, Facebook, DoubleClick.',
     'High'],
    ['Personalization (TrenDemon)', 'Content recommendation engine; evaluate EDS-compatible alternatives or direct integration.',
     'High'],
    ['Authentication (My NEC)', 'Event registration flows via auth.nec.com; ensure proper URL/redirect handling.',
     'Medium'],
    ['Contact Form Replacement', 'Replace legacy CGI contact form with modern alternative.',
     'Low'],
    ['Content Migration (Bulk)', 'Migrate ~1,362 pages of Japanese content with proper encoding, images, and metadata.',
     'High (volume)'],
    ['Footer', 'Comprehensive mega footer with 60+ categorized links.',
     'Medium'],
]
add_styled_table(doc,
    ['Work Stream', 'Description', 'Effort Level'],
    effort,
    col_widths=[1.8, 3.5, 1.0])

doc.add_heading('8.4 Key Migration Risks', level=2)
risks = [
    ['Japanese Text Rendering', 'Medium', 'Japanese typography (line-breaking, mixed-width characters) must render correctly across browsers'],
    ['Dual TMS Complexity', 'High', 'Running both Tealium and GTM requires careful parallel implementation to avoid tag conflicts'],
    ['Dynamic Content Filtering', 'High', 'Current JS-based filter system needs EDS-compatible redesign; may impact SEO'],
    ['Content Volume', 'Medium', 'Migrating ~1,362 pages requires robust import scripting with Japanese encoding handling'],
    ['Authentication Integration', 'Medium', 'My NEC SSO redirects with Tealium tracking parameters must be preserved'],
    ['TrenDemon Personalization', 'High', 'Content personalization behavior may differ on EDS; needs testing and possible alternative'],
    ['Video Platform Dependency', 'Medium', 'Brightcove player embeds with custom Tealium plugin need validation on EDS'],
    ['SEO Impact', 'High', 'Japanese SEO requires careful URL mapping, canonical handling, and sitemap generation'],
]
add_styled_table(doc,
    ['Risk Area', 'Severity', 'Description'],
    risks,
    col_widths=[1.5, 0.7, 4.1])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 9. SCREENSHOTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Screenshots', level=1)

screenshots = [
    ('nec-wisdom-homepage.png', 'Figure 1: Homepage - Hero carousel, category navigation, latest articles'),
    ('nec-wisdom-feature-listing.png', 'Figure 2: Theme/Topic Listing - AI articles with card grid and "Load More"'),
    ('nec-wisdom-special.png', 'Figure 3: Special Feature Collection - Digital Ethics curated editorial'),
    ('nec-wisdom-series.png', 'Figure 4: Series/Author Column - North America Trends by Koichi Oda'),
    ('nec-wisdom-event.png', 'Figure 5: Event Hub - Featured events, event reports, and past events'),
    ('nec-wisdom-event-detail.png', 'Figure 6: Event Detail - Seminar with video player, sessions, and speaker profiles'),
]

for path, caption in screenshots:
    full_path = f'/workspace/{path}'
    if os.path.exists(full_path):
        add_screenshot(doc, full_path, caption, width=5.0)
        doc.add_paragraph()  # spacing
    else:
        doc.add_paragraph(f'[Screenshot not available: {path}]')

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_path = '/workspace/NEC_Wisdom_Site_Analysis_Report.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
print(f'File size: {os.path.getsize(output_path):,} bytes')
