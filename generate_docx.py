#!/usr/bin/env python3
"""Generate a professionally formatted MS Word document for the Atlassian.com site analysis."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Definitions ───────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)  # Atlassian blue

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# Helper: shaded cell
BLUE_BG = 'D9E2F3'
DARK_BLUE_BG = '0052CC'
LIGHT_GREY = 'F2F2F2'

def set_cell_bg(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, size=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = 'Calibri'
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with styled header row and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, DARK_BLUE_BG)
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=Pt(9))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_bg(cell, LIGHT_GREY)
            set_cell_text(cell, val, size=Pt(9))

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Atlassian.com')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)
run.bold = True
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Website Analysis Report')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
run.font.name = 'Calibri'

doc.add_paragraph('')

line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = line.add_run('━' * 60)
run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)
run.font.size = Pt(10)

doc.add_paragraph('')

details = [
    ('Source URL:', 'https://www.atlassian.com/'),
    ('Date:', 'March 2, 2026'),
    ('Purpose:', 'Pre-migration site analysis for Adobe Edge Delivery Services'),
    ('Prepared by:', 'Automated Site Analysis'),
]
for label, value in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Templates Inventory',
    '2. Blocks / Components Catalog',
    '   2.1 Global Components',
    '   2.2 Content Blocks',
    '3. Page Counts by Template',
    '   3.1 Detailed Breakdown',
    '   3.2 Migration Type Summary',
    '4. Integrations Analysis',
    '5. Complex Use Cases & Observations',
    '6. Migration Estimates',
    '   6.1 Effort Breakdown by Phase',
    '   6.2 Summary Estimates',
    '   6.3 Team Composition Recommendation',
    '   6.4 Key Risks & Considerations',
    'Appendix: Screenshots',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. TEMPLATES INVENTORY
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Templates Inventory', level=1)
doc.add_paragraph(
    'The following 20 unique page templates were identified across the Atlassian.com website. '
    'Each template represents a distinct page layout pattern used for a specific content purpose.'
)

templates = [
    ['1', 'Homepage', 'High', 'Complex hero with video, animated text, multiple interactive sections (tabs, carousels), product showcase cards with testimonials, statistics counter, customer logo carousel', 'https://www.atlassian.com/'],
    ['2', 'Product Page', 'High', 'Inline signup form, animated hero, video embeds, feature showcase with images, horizontal scrolling card lists, customer testimonial, integration showcase, Gartner badge section', 'https://www.atlassian.com/software/jira\nhttps://www.atlassian.com/software/confluence\nhttps://www.atlassian.com/software/loom'],
    ['3', 'Product Pricing', 'High', 'Dynamic pricing calculator, tier comparison table with toggle (monthly/annual), feature comparison matrix with checkmarks, FAQ accordion, interactive user count slider', 'https://www.atlassian.com/software/jira/pricing'],
    ['4', 'Customer Story', 'Medium', 'Hero with quote/avatar, key results metrics, company info sidebar, long-form article body with blockquotes, embedded video, related content cards at bottom', 'https://www.atlassian.com/customers/lumen\nhttps://www.atlassian.com/customers/dropbox'],
    ['5', 'Customer Stories Listing', 'Medium', 'Filterable grid of customer cards with logos, industry/size filtering, search, pagination', 'https://www.atlassian.com/customers'],
    ['6', 'Blog Listing', 'Medium', 'Featured hero article, 4-article grid, latest stories list with images, featured collections carousel, newsletter signup CTA, "Load more" pagination', 'https://www.atlassian.com/blog'],
    ['7', 'Blog Article', 'Medium', 'Long-form content with images, author/category metadata, inline CTAs, social sharing, related articles, newsletter signup', 'https://www.atlassian.com/blog/teamwork/work-needs-a-why-7-team-practices-for-sharing-purpose'],
    ['8', 'Solutions / Teams Page', 'Medium', 'Hero section, feature blocks with illustrations, tabbed content areas, template cards, customer testimonials', 'https://www.atlassian.com/teams/software-development\nhttps://www.atlassian.com/teams/marketing'],
    ['9', 'Collection / Bundle Page', 'Medium', 'Product bundle hero, product comparison cards, feature highlights, pricing overview, cross-sell blocks', 'https://www.atlassian.com/collections/teamwork\nhttps://www.atlassian.com/collections/strategy'],
    ['10', 'Industry Page', 'Medium', 'Industry-specific hero, solution mapping, customer logos from industry, case study highlights', 'https://www.atlassian.com/industries/retail\nhttps://www.atlassian.com/industries/telecom'],
    ['11', 'Features Page', 'Medium', 'Feature grid with icons/descriptions, feature detail sections with screenshots, comparison tables', 'https://www.atlassian.com/software/jira/features'],
    ['12', 'Template Gallery Page', 'Low', 'Template card with description, screenshot preview, "Try it" CTA, related templates', 'https://www.atlassian.com/software/jira/templates/scrum'],
    ['13', 'Enterprise Page', 'High', 'Custom hero, interactive ROI calculator, enterprise-specific feature cards, security/compliance badges, contact form', 'https://www.atlassian.com/enterprise'],
    ['14', 'Company / About Page', 'Medium', 'Mission/values hero, team photos, office locations, culture content', 'https://www.atlassian.com/company'],
    ['15', 'Careers Page', 'High', 'Job search with filters, location-based listings, department categorization, benefits section, culture content', 'https://www.atlassian.com/company/careers'],
    ['16', 'Trust / Security Page', 'Medium', 'Compliance badges, security certifications grid, documentation links, status page embeds', 'https://www.atlassian.com/trust'],
    ['17', 'Legal / Policy Page', 'Low', 'Long-form text content, table of contents sidebar, section anchors', 'https://www.atlassian.com/legal/privacy-policy'],
    ['18', 'Contact Page', 'Medium', 'Contact form, regional office info, support links, partner finder', 'https://www.atlassian.com/company/contact'],
    ['19', 'Migration / Cloud Page', 'Medium', 'Step-by-step migration guide, assessment tools, resource links, comparison tables', 'https://www.atlassian.com/migration/cloud'],
    ['20', 'Resources Landing Page', 'Medium', 'Categorized resource cards, filtering by type/topic, search', 'https://www.atlassian.com/resources'],
]

add_styled_table(doc,
    ['#', 'Template Name', 'Complexity', 'Reasoning', 'Reference URL(s)'],
    templates,
    col_widths=[0.8, 3, 1.5, 6.5, 5.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. BLOCKS / COMPONENTS CATALOG
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Blocks / Components Catalog', level=1)
doc.add_paragraph(
    '32 reusable blocks/components were identified across the site. Design variations of the same '
    'content model are tracked as variants of a single block rather than separate blocks.'
)

# 2.1 Global Components
doc.add_heading('2.1 Global Components', level=2)

global_blocks = [
    ['1', 'Global Header / Navigation', 'High', 'Mega-menu navigation with multi-level flyouts organized by Products (by role), Solutions (by Use Case, Team, Size, Industry), Why Atlassian, Resources, and Enterprise. Includes search, sign-in, and CTA button. Responsive hamburger menu on mobile.', 'All pages'],
    ['2', 'Global Footer', 'Medium', '4-column layout with Company links, Products, Resources, and Learn sections. Includes copyright, privacy links, language selector dropdown (12 languages), and social proof elements.', 'All pages'],
    ['3', 'Cookie Consent Banner', 'Low', 'OneTrust-powered cookie banner with "Manage preferences" and "I understand" buttons. Links to Privacy Policy and Global Privacy Control.', 'All pages'],
    ['4', 'Search Overlay', 'Medium', 'Full-screen search overlay triggered by search icon in header. Powers global site search.', 'All pages'],
]

add_styled_table(doc,
    ['#', 'Block Name', 'Complexity', 'Description', 'Reference URL(s)'],
    global_blocks,
    col_widths=[0.8, 3.5, 1.5, 7.5, 3]
)

doc.add_paragraph('')

# 2.2 Content Blocks
doc.add_heading('2.2 Content Blocks', level=2)

content_blocks = [
    ['5', 'Hero Banner (Video)', 'High', 'Full-width hero with animated headline text, subtext, CTA buttons, and background video with play button overlay. Multiple design variations: event takeover, product launch, announcement.', 'https://www.atlassian.com/'],
    ['6', 'Hero Banner (Signup Form)', 'High', 'Hero with inline email signup form, Google SSO button, and product category cards below. Used on product pages.', 'https://www.atlassian.com/software/jira'],
    ['7', 'Hero Banner (Text + Image)', 'Medium', 'Hero with heading, description text, CTA button(s), and side/background image. Standard variant used across solutions and feature pages.', 'Solutions, Features pages'],
    ['8', 'Hero Banner (Customer Story)', 'Medium', 'Hero with large heading, featured quote with avatar image, author name/title. Yellow accent bar at top.', 'https://www.atlassian.com/customers/lumen'],
    ['9', 'Product Cards Grid', 'Medium', 'Horizontal card layout showcasing products (Jira, Confluence, Loom, etc.). Each card has product icon, heading, description, "Get it free" and "Explore" CTAs, plus a customer testimonial.', 'https://www.atlassian.com/'],
    ['10', 'Feature Cards (Horizontal Scroll)', 'Medium', 'Horizontally scrollable list of feature cards, each with illustration, heading, description, and link. Used for showcasing capabilities.', 'https://www.atlassian.com/software/jira'],
    ['11', 'Feature Section (Text + Media)', 'Medium', 'Split layout with text content (heading, description, CTA link) on one side and image/video on the other. Multiple orientation variations (left/right).', 'https://www.atlassian.com/software/jira'],
    ['12', 'Statistics / Metrics Counter', 'Medium', 'Animated number counters with descriptions (e.g., "300,000+ companies", "200+ countries", "80% of Fortune 500"). Used for social proof.', 'https://www.atlassian.com/'],
    ['13', 'Customer Logo Bar', 'Low', 'Horizontal strip of grayscale customer logos (NASA, Visa, Canva, Ford, etc.). Used for social proof.', 'https://www.atlassian.com/software/jira'],
    ['14', 'Tabbed Content Section', 'High', 'Tab navigation (Software, Product Management, Marketing, etc.) with different content panels. Each panel has heading, description, link, and template cards below.', 'https://www.atlassian.com/'],
    ['15', 'Customer Testimonial Carousel', 'High', 'Logo-tabbed carousel of customer quotes. Tab bar shows customer logos; content shows company size tag, blockquote, author info, and case study link. 14 testimonials rotating.', 'https://www.atlassian.com/'],
    ['16', 'Video Embed Block', 'Medium', 'Vimeo video player with custom play button overlay and thumbnail. Used inline and in hero sections.', 'https://www.atlassian.com/'],
    ['17', 'Announcement Banner', 'Low', 'Dismissible top banner with text message and "Learn More" link. Used for promotions (e.g., Gartner recognition).', 'https://www.atlassian.com/software/jira'],
    ['18', 'CTA Section (Full-width)', 'Low', 'Full-width colored background section with heading, description, and CTA button. Used as page-closing call to action.', 'https://www.atlassian.com/'],
    ['19', 'Blockquote / Pull Quote', 'Low', 'Styled blockquote with author name, title, and optional avatar image. Used in customer stories and product pages.', 'https://www.atlassian.com/customers/lumen'],
    ['20', 'Key Results / Metrics Block', 'Low', 'Definition list of key metrics (e.g., "200% increase in throughput") with descriptions. Used in customer stories.', 'https://www.atlassian.com/customers/lumen'],
    ['21', 'Company Info Sidebar', 'Low', 'Sidebar card with company logo, industry, user count, location, and links to full story/PDF download.', 'https://www.atlassian.com/customers/lumen'],
    ['22', 'Related Content Cards', 'Low', '3-column grid of related content cards with heading, description, and CTA link. Used at bottom of detail pages.', 'https://www.atlassian.com/customers/lumen'],
    ['23', 'Blog Article Card', 'Low', 'Card with thumbnail image, heading, excerpt, author name, and category tag. Two variants: list view and grid view.', 'https://www.atlassian.com/blog'],
    ['24', 'Blog Featured Hero', 'Medium', 'Large featured article card with full-width image, heading, excerpt, author, and category. Used at top of blog listing.', 'https://www.atlassian.com/blog'],
    ['25', 'Newsletter Signup Block', 'Medium', 'Inline newsletter subscription with heading, description, and CTA button. Two design variants (inline and footer).', 'https://www.atlassian.com/blog'],
    ['26', 'Collection Card (Blog)', 'Low', 'Card with collection name badge, heading, description, and "View Collection" link. Used in horizontal carousel.', 'https://www.atlassian.com/blog'],
    ['27', 'Pricing Tier Cards', 'High', 'Pricing comparison cards with tier name, price, user count, feature list, and CTA. Includes toggle for monthly/annual pricing and interactive user slider.', 'https://www.atlassian.com/software/jira/pricing'],
    ['28', 'Feature Comparison Table', 'High', 'Multi-column comparison matrix with checkmarks/crosses for feature availability across tiers. Expandable category sections.', 'https://www.atlassian.com/software/jira/pricing'],
    ['29', 'FAQ Accordion', 'Low', 'Expandable accordion sections for frequently asked questions.', 'Pricing, Product pages'],
    ['30', 'Gartner Badge Section', 'Low', 'Gartner Magic Quadrant recognition badge with description and links to reports.', 'https://www.atlassian.com/software/jira'],
    ['31', 'Product Collection Banner', 'Medium', 'Banner promoting product bundles (Teamwork Collection) with product icons, heading, description, dual CTAs, and illustration.', 'https://www.atlassian.com/'],
    ['32', 'Event Takeover Section', 'High', 'Full-width event promotion with event logo, date/location info with icons, heading, description, dual CTAs, and background video. Time-limited promotional content.', 'https://www.atlassian.com/'],
]

add_styled_table(doc,
    ['#', 'Block Name', 'Complexity', 'Description', 'Reference URL(s)'],
    content_blocks,
    col_widths=[0.8, 3.5, 1.5, 7.5, 3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. PAGE COUNTS BY TEMPLATE
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Page Counts by Template', level=1)

doc.add_heading('3.1 Detailed Breakdown', level=2)

page_counts = [
    ['Homepage', '1 (+12 localized)', 'Manual', 'Complex interactive elements, animations, event takeover'],
    ['Product Pages', '~30', 'Semi-automated', 'Shared layout but unique content/media per product. ~156 URLs in /software/* including sub-pages'],
    ['Product Pricing', '~8', 'Manual', 'Dynamic pricing calculator, tier logic, custom per product'],
    ['Product Features', '~10', 'Semi-automated', 'Feature grids vary per product'],
    ['Template Gallery', '~20', 'Automated', 'Standardized card structure'],
    ['Customer Stories', '~18', 'Automated', 'Consistent template with variable long-form content'],
    ['Customer Listing', '1', 'Manual', 'Filtering/search functionality'],
    ['Blog Listing', '~8', 'Semi-automated', 'Category-specific listings (teamwork, productivity, etc.)'],
    ['Blog Articles', '~500+', 'Automated', 'Standard long-form content (not in sitemap, dynamic CMS)'],
    ['Solutions / Teams', '~12', 'Semi-automated', 'Shared structure, unique content'],
    ['Collection / Bundle', '4', 'Semi-automated', 'Standard layout per bundle'],
    ['Industry Pages', '3+', 'Semi-automated', 'Standard layout per industry'],
    ['Enterprise', '~5', 'Manual', 'Custom interactive elements'],
    ['Company / About', '~6', 'Semi-automated', 'Various company info pages'],
    ['Careers', '1+', 'Manual', 'Job search with dynamic data integration'],
    ['Trust / Security', '~3', 'Semi-automated', 'Compliance/certification content'],
    ['Legal / Policy', '~5', 'Automated', 'Long-form text, minimal interactivity'],
    ['Contact', '1', 'Manual', 'Form integration required'],
    ['Migration / Cloud', '~5', 'Semi-automated', 'Assessment tools, custom widgets'],
    ['Resources Landing', '~3', 'Semi-automated', 'Filterable resource hub'],
    ['Localized Variants', '~120', 'Automated', '10+ languages across key pages'],
]

add_styled_table(doc,
    ['Template', 'Estimated Count', 'Migration Type', 'Notes'],
    page_counts,
    col_widths=[3.5, 2.5, 2.5, 8]
)

doc.add_paragraph('')

p = doc.add_paragraph()
r = p.add_run('Total estimated pages in sitemap: ')
r.bold = True
p.add_run('~234 (English) + localized variants')
p = doc.add_paragraph()
r = p.add_run('Total estimated pages including blog/dynamic: ')
r.bold = True
p.add_run('~800-1,000+')

doc.add_paragraph('')
doc.add_heading('3.2 Migration Type Summary', level=2)

migration_summary = [
    ['Automated', '~550-650 pages', 'Blog articles, customer stories, legal pages, template gallery, localized variants'],
    ['Semi-automated', '~200-250 pages', 'Product pages, solutions, features, teams, industry, collections'],
    ['Manual', '~50-100 pages', 'Homepage, pricing, enterprise, careers, contact, customer listing'],
]

add_styled_table(doc,
    ['Migration Type', 'Estimated Page Count', 'Description'],
    migration_summary,
    col_widths=[3, 3, 10.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. INTEGRATIONS ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Integrations Analysis', level=1)
doc.add_paragraph(
    '21 third-party integrations and embedded services were identified through network request '
    'analysis, script tag inspection, and iframe detection.'
)

integrations = [
    ['1', 'Google Tag Manager (GTM)', 'Tag management / embed', 'Low', 'Tag container for analytics and marketing pixels', 'All pages'],
    ['2', 'Google Analytics (GA4)', 'API / embed', 'Low', 'Web analytics and conversion tracking', 'All pages'],
    ['3', 'Google reCAPTCHA', 'Embed / API', 'Medium', 'Bot protection on forms and signup flows', 'All pages (iframe)'],
    ['4', 'OneTrust (CookieLaw)', 'Embed / custom code', 'Medium', 'Cookie consent management, GDPR/CCPA compliance', 'All pages'],
    ['5', 'Vimeo Player', 'Embed', 'Low', 'Video hosting and playback for product demos', 'Homepage, Product pages'],
    ['6', 'Marketo (Munchkin)', 'API / embed', 'High', 'Marketing automation, lead capture, form handling', 'All pages'],
    ['7', 'Facebook Pixel', 'Embed', 'Low', 'Advertising conversion tracking and retargeting', 'All pages'],
    ['8', 'Twitter/X Ads', 'Embed', 'Low', 'Advertising conversion tracking', 'All pages'],
    ['9', 'LinkedIn Insight Tag', 'Embed', 'Low', 'B2B advertising and conversion tracking', 'All pages'],
    ['10', 'Bing / Microsoft Ads', 'Embed', 'Low', 'Search advertising conversion tracking', 'All pages'],
    ['11', 'Reddit Pixel', 'Embed', 'Low', 'Advertising conversion tracking', 'All pages'],
    ['12', '6sense', 'API / embed', 'Medium', 'ABM intent data, visitor identification', 'All pages'],
    ['13', 'TechTarget / Priority Engine', 'API / embed', 'Medium', 'B2B intent data and lead enrichment', 'All pages'],
    ['14', 'Impact.com', 'Embed', 'Medium', 'Affiliate/partner tracking', 'All pages'],
    ['15', 'Google DoubleClick (DV360)', 'Embed', 'Low', 'Programmatic advertising, Floodlight tags', 'All pages'],
    ['16', 'ClickTale / Contentsquare', 'Embed', 'Medium', 'Session recording and heatmap analytics', 'All pages'],
    ['17', 'Intercom', 'Embed / API', 'High', 'Customer support chat widget (launcher disabled but SDK loaded)', 'Product & Customer pages'],
    ['18', 'Sentry', 'API', 'Medium', 'Error monitoring and application performance', 'Main WAC app (React)'],
    ['19', 'Atlassian Identity (Guard)', 'Custom API', 'High', 'SSO authentication, user management, account linking', 'Sign-in flows'],
    ['20', 'WordPress (Blog)', 'CMS / custom', 'High', 'Blog CMS on separate stack (jQuery Migrate, different header/footer)', '/blog/* pages only'],
    ['21', 'Contentful', 'CMS / API', 'High', 'Content management for customer stories (assets.ctfassets.net)', 'Customer stories'],
]

add_styled_table(doc,
    ['#', 'Integration Name', 'Type', 'Complexity', 'Purpose', 'Pages Used'],
    integrations,
    col_widths=[0.8, 3.5, 2.5, 1.5, 4.5, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. COMPLEX USE CASES & OBSERVATIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Complex Use Cases & Observations', level=1)
doc.add_paragraph(
    'The following complex behaviors, edge cases, and functionality require special attention during migration.'
)

complex_cases = [
    ['1', 'Dual CMS Architecture', '2 systems', 'Blog (/blog/*) vs. main site', 'Blog runs on WordPress; main site is a React SPA (WAC). Different headers, footers, JS stacks, and routing. Requires two separate migration strategies.'],
    ['2', 'Dynamic Pricing Calculator', '~8 pages', '/software/*/pricing', 'Interactive pricing with user count slider, monthly/annual toggle, tier comparison. Server-side pricing API. Cannot be statically migrated.'],
    ['3', 'Inline Signup / Authentication', '~30+ pages', 'Product pages, CTAs', 'Email signup forms with Google SSO, connected to Atlassian Identity (Guard) and account provisioning API. Real-time validation and redirect flows.'],
    ['4', 'Multi-language / i18n', '12 languages', '~120 localized pages', 'Full URL-based localization (/de/, /ja/, /fr/, etc.) with language selector. Content translation and locale-specific pricing/legal differences.'],
    ['5', 'Event Takeover (Time-bound)', '1 active', 'Homepage', 'Temporary hero section promoting Team \'26 conference. Overrides normal homepage hero. Likely managed via feature flags or CMS scheduling.'],
    ['6', 'Customer Testimonial Carousel', '2-3 pages', 'Homepage, Product pages', '14-item logo-tabbed carousel with animation transitions. Complex interaction state management.'],
    ['7', 'Animated Text / Counter Effects', '5-10 pages', 'Homepage, Product pages', 'Animated headline text with cursor animation, number counter animations with scroll-triggered start, SVG animations.'],
    ['8', 'Video Embeds with Custom Controls', '10+ pages', 'Homepage, Product, Customer Stories', 'Vimeo videos with custom play button overlay, lazy-loaded iframes, custom thumbnail images.'],
    ['9', 'Mega-Menu Navigation', 'Global', 'All pages', 'Multi-level flyout navigation organized by Products (6 sub-categories), Solutions (4 groupings), Resources, Why Atlassian. Complex responsive behavior.'],
    ['10', 'React SPA Routing', 'Main site', 'All WAC pages', 'Main site (excluding blog) is a React SPA with client-side routing. Affects SEO strategy and page loading patterns during migration.'],
    ['11', 'ABM & Intent Data Layer', 'All pages', 'Global', '6sense and TechTarget integration create sophisticated visitor identification and personalization. Content may be dynamically adjusted based on visitor data.'],
    ['12', 'Intercom Chat Integration', 'Product pages', 'Product / Customer pages', 'Chat SDK loaded but launcher disabled via settings - indicates conditional enablement, possibly based on user state or A/B testing.'],
    ['13', 'PDF Downloads', '~18 pages', 'Customer stories', 'Case study PDFs hosted on Contentful CDN. Require asset migration and link management.'],
    ['14', 'Feature Flag System', 'Global', 'All pages', 'Multiple FeatureGateClient instances detected. Content and features may be conditionally rendered, complicating static migration.'],
]

add_styled_table(doc,
    ['#', 'Complex Behavior', 'Instances', 'Location', 'Why It\'s Complex'],
    complex_cases,
    col_widths=[0.8, 3.5, 1.8, 3, 7.2]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. MIGRATION ESTIMATES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Migration Estimates', level=1)

doc.add_heading('6.1 Effort Breakdown by Phase', level=2)

phases = [
    ['1. Discovery & Planning', 'Site architecture, content inventory, template mapping', '3-4 weeks', 'Template finalization, content audit, integration mapping'],
    ['2. Design System Migration', 'CSS custom properties, typography, colors, spacing, responsive breakpoints', '2-3 weeks', 'Extract from WAC styles, map to Edge Delivery Services tokens'],
    ['3. Block Development', '32 identified blocks / components', '6-8 weeks', '~2 blocks/week avg. High-complexity blocks (pricing, mega-nav, carousels) take 1 week each'],
    ['4. Template Development', '20 templates', '4-5 weeks', 'Reuses blocks; mostly layout/composition work'],
    ['5. Content Migration - Automated', '~550-650 pages (blog articles, legal, customer stories, template gallery)', '2-3 weeks', 'Scripted import with manual QA sampling'],
    ['6. Content Migration - Semi-automated', '~200-250 pages (product, solutions, teams, features)', '4-5 weeks', 'Template-based import with manual content adjustment'],
    ['7. Content Migration - Manual', '~50-100 pages (homepage, pricing, enterprise, careers, contact)', '4-6 weeks', 'Custom interactive elements, form integrations, dynamic content'],
    ['8. Integration Reimplementation', '21 integrations identified', '3-4 weeks', 'Analytics/marketing pixels (1 wk), forms/Marketo (1-2 wks), auth flows (1 wk)'],
    ['9. Localization', '~120 localized page variants across 12 languages', '2-3 weeks', 'Automated with translation management, URL structure setup'],
    ['10. Blog Migration', 'WordPress to Edge Delivery Services', '3-4 weeks', 'Separate CMS migration, ~500+ articles, different template system'],
    ['11. QA & Testing', 'Full regression, cross-browser, accessibility, performance', '4-5 weeks', 'Parallel with content migration phases'],
    ['12. Performance Optimization', 'Core Web Vitals, Lighthouse scoring', '1-2 weeks', 'Final tuning after migration'],
]

add_styled_table(doc,
    ['Phase', 'Scope', 'Effort Estimate', 'Notes'],
    phases,
    col_widths=[3.5, 5, 2, 6]
)

doc.add_paragraph('')
doc.add_heading('6.2 Summary Estimates', level=2)

summary = [
    ['Automated migration', '4-6 weeks', 'Weeks 8-13'],
    ['Semi-automated migration', '4-5 weeks', 'Weeks 8-13'],
    ['Manual / custom migration', '4-6 weeks', 'Weeks 10-16'],
    ['Block & template development', '10-13 weeks', 'Weeks 2-14'],
    ['Integration work', '3-4 weeks', 'Weeks 12-16'],
    ['QA & testing', '4-5 weeks', 'Weeks 12-18'],
    ['Total project duration', '~18-22 weeks (parallel)', ''],
    ['Total effort', '~35-50 person-weeks', ''],
]

add_styled_table(doc,
    ['Category', 'Effort', 'Duration (parallel)'],
    summary,
    col_widths=[5, 4, 4]
)

doc.add_paragraph('')
doc.add_heading('6.3 Team Composition Recommendation', level=2)

team = [
    ['Project Lead / Architect', '1', 'Full duration'],
    ['Edge Delivery Services Developer (blocks/templates)', '2', '14 weeks'],
    ['Content Migration Engineer', '1-2', '10 weeks'],
    ['Frontend Developer (integrations, interactive)', '1', '8 weeks'],
    ['QA Engineer', '1', '8 weeks'],
    ['Total team size', '6-7', ''],
]

add_styled_table(doc,
    ['Role', 'FTEs', 'Duration'],
    team,
    col_widths=[7, 2, 3]
)

doc.add_paragraph('')
doc.add_heading('6.4 Key Risks & Considerations', level=2)

risks = [
    ('Dual CMS (WordPress Blog)',
     'The blog runs on a completely separate WordPress stack and requires its own migration strategy. This is the single largest content volume.'),
    ('React SPA Architecture',
     'The main site is a React SPA (WAC), not traditional server-rendered pages. Scraping and content extraction will require JavaScript rendering.'),
    ('Dynamic Pricing',
     'Pricing calculators with real-time API calls cannot be statically migrated and will need custom Edge Delivery Services blocks with API integration.'),
    ('Feature Flags',
     'Content may render differently based on feature flag state, which could complicate content capture during migration.'),
    ('ABM Personalization',
     'If 6sense/TechTarget powers content personalization, some page variants may not be visible during standard crawling.'),
    ('Atlassian Identity (Auth)',
     'Signup and sign-in flows connect to Atlassian\'s identity platform and will need to be maintained as external integrations.'),
]

for i, (title, desc) in enumerate(risks, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'{i}. {title}: ')
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(desc).font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX: SCREENSHOTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading('Appendix: Screenshots', level=1)
doc.add_paragraph(
    'The following screenshots were captured during the analysis for visual reference.'
)

screenshots = [
    ('atlassian-homepage.png', 'Homepage — https://www.atlassian.com/'),
    ('atlassian-jira-product.png', 'Jira Product Page — https://www.atlassian.com/software/jira'),
    ('atlassian-blog-listing.png', 'Blog Listing — https://www.atlassian.com/blog'),
    ('atlassian-jira-pricing.png', 'Jira Pricing — https://www.atlassian.com/software/jira/pricing'),
    ('atlassian-customer-story.png', 'Customer Story (Lumen) — https://www.atlassian.com/customers/lumen'),
]

for fname, caption in screenshots:
    fpath = os.path.join('/workspace', fname)
    if os.path.exists(fpath):
        doc.add_heading(caption, level=3)
        try:
            doc.add_picture(fpath, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[Screenshot: {fname} — could not embed: {e}]')
        doc.add_paragraph('')
    else:
        doc.add_paragraph(f'[Screenshot not found: {fname}]')

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_path = '/workspace/Atlassian_Site_Analysis_Report.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
